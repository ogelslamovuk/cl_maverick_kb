"""Normalize local inbox materials into Obsidian-friendly markdown files.

Supported input formats: TXT, MD, DOCX, PDF.
No source files are deleted or moved.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}
TEXT_EXTENSIONS = {".txt", ".md"}
ENCODINGS = ("utf-8-sig", "utf-8", "cp1251")


def project_root() -> Path:
    return Path(__file__).resolve().parents[1]


def yaml_quote(value: str) -> str:
    escaped = value.replace("\\", "\\\\").replace('"', '\\"')
    return f'"{escaped}"'


def read_text_file(path: Path) -> str:
    last_error: Exception | None = None
    for encoding in ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc
    raise ValueError(f"could not decode text file: {last_error}")


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("missing dependency: python-docx") from exc

    document = Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("missing dependency: pypdf") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"<!-- page: {index} -->\n{text.strip()}")
    return "\n\n".join(pages)


def extract_text(path: Path) -> str:
    extension = path.suffix.lower()
    if extension in TEXT_EXTENSIONS:
        return read_text_file(path)
    if extension == ".docx":
        return extract_docx_text(path)
    if extension == ".pdf":
        return extract_pdf_text(path)
    raise ValueError(f"unsupported file type: {extension or '<none>'}")


def build_output_markdown(
    source_path: Path,
    source_type: str,
    processed_at: str,
    extracted_text: str,
) -> str:
    return "\n".join(
        [
            "---",
            f"source_file: {yaml_quote(source_path.name)}",
            f"source_type: {yaml_quote(source_type)}",
            f"processed_at: {yaml_quote(processed_at)}",
            "status: raw_normalized",
            "---",
            "",
            f"# Raw normalized: {source_path.name}",
            "",
            "## Extracted text",
            "",
            extracted_text.rstrip(),
            "",
        ]
    )


def append_log(
    log_path: Path,
    processed_at: str,
    source_name: str,
    source_type: str,
    result_path: Path | None,
    status: str,
    error: str = "",
) -> None:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    if not log_path.exists():
        log_path.write_text("# Processing Log\n\n", encoding="utf-8")

    result_value = str(result_path) if result_path else "-"
    error_value = error if error else "-"
    entry = "\n".join(
        [
            f"## {processed_at}",
            "",
            f"- source_file: {source_name}",
            f"- source_type: {source_type}",
            f"- result_path: {result_value}",
            f"- status: {status}",
            f"- error: {error_value}",
            "",
        ]
    )
    with log_path.open("a", encoding="utf-8", newline="\n") as handle:
        handle.write(entry)


def output_path_for(source_path: Path, output_dir: Path, date_prefix: str) -> Path:
    return output_dir / f"{date_prefix}__{source_path.name}.md"


def process_file(source_path: Path, output_dir: Path, log_path: Path) -> None:
    now = datetime.now().astimezone()
    processed_at = now.isoformat(timespec="seconds")
    date_prefix = now.strftime("%Y-%m-%d")
    source_type = source_path.suffix.lower().lstrip(".") or "unknown"

    if source_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        append_log(
            log_path=log_path,
            processed_at=processed_at,
            source_name=source_path.name,
            source_type=source_type,
            result_path=None,
            status="error",
            error=f"unsupported file type: {source_path.suffix or '<none>'}",
        )
        return

    result_path = output_path_for(source_path, output_dir, date_prefix)

    try:
        extracted_text = extract_text(source_path)
        if not extracted_text.strip():
            raise ValueError("no text extracted")

        output_dir.mkdir(parents=True, exist_ok=True)
        result_path.write_text(
            build_output_markdown(
                source_path=source_path,
                source_type=source_type,
                processed_at=processed_at,
                extracted_text=extracted_text,
            ),
            encoding="utf-8",
            newline="\n",
        )
        append_log(
            log_path=log_path,
            processed_at=processed_at,
            source_name=source_path.name,
            source_type=source_type,
            result_path=result_path,
            status="success",
        )
    except Exception as exc:  # Keep processing the rest of the inbox.
        append_log(
            log_path=log_path,
            processed_at=processed_at,
            source_name=source_path.name,
            source_type=source_type,
            result_path=result_path,
            status="error",
            error=str(exc),
        )


def main() -> int:
    root = project_root()
    inbox_dir = root / "inbox"
    output_dir = root / "Maverick_KB" / "raw_normalized"
    log_path = root / "Maverick_KB" / "PROCESSING_LOG.md"

    if not inbox_dir.exists():
        raise SystemExit(f"Inbox directory not found: {inbox_dir}")

    for source_path in sorted(inbox_dir.iterdir(), key=lambda item: item.name.lower()):
        if not source_path.is_file():
            continue
        if source_path.name.lower() == "readme.md":
            continue
        process_file(source_path=source_path, output_dir=output_dir, log_path=log_path)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
